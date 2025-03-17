import base64
import requests
from docx import Document

# Define API URL
API_URL = "http://localhost:8000/upload_file"

# Create a sample .docx file
doc = Document()
doc.add_heading("Sample Document", level=1)
doc.add_paragraph("This is a test document for verifying SharePoint file upload via API.")
doc_path = "sample_test.docx"
doc.save(doc_path)

# Read and encode file content as base64
with open(doc_path, "rb") as file:
    encoded_file = base64.b64encode(file.read()).decode("utf-8")

# Define request payload
payload = {
    "file_name": "test_document.docx",  # Change to desired file name
    "server_id": "e8ad33bd-c780-4630-a6ce-9a93c83a9480",   # Replace with actual folder ID in SharePoint
    "file_data": encoded_file
}

# Send POST request
response = requests.post(API_URL, json=payload)

# Print response
print("Status Code:", response.status_code)
print("Response:", response.json())